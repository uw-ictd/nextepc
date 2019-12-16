# Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>

# This file is part of Open5GS.

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (C) 2019 by Sukchan Lee <acetcom@gmail.com>
 *
 * This file is part of Open5GS.
 *
 * This program is free software: you can redistribute it and/or modify
 * it under the terms of the GNU Affero General Public License as published by
 * the Free Software Foundation, either version 3 of the License, or
 * (at your option) any later version.
 *
 * This program is distributed in the hope that it will be useful,
 * but WITHOUT ANY WARRANTY; without even the implied warranty of
 * MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
 * GNU General Public License for more details.
 *
 * You should have received a copy of the GNU General Public License
 * along with this program.  If not, see <https://www.gnu.org/licenses/>.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by pfcp-tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print "Python generating TLV build/parser for PFCP v%s" % (version)
    print "Usage: python pfcp-tlv.py [options]"
    print "Available options:"
    print "-d        Enable script debug"
    print "-f [file] Input file to parse"
    print "-o [dir]  Output files to given directory"
    print "-c [dir]  Cache files to given directory"
    print "-h        Print this help and return"

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def get_cells(cells):
    #instance = cells[4].text.encode('ascii', 'ignore')
    #if instance.isdigit() is not True:
    #    return None
    instance = "0"  # PFCP has no instance
    note = cells[0].text.encode('ascii', 'ignore')
    if note.find('NOTE') != -1:
        return None
    comment = cells[2].text.encode('ascii', 'ignore')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);
    #print comment
    ie_type = re.sub('\s*$', '', re.sub('\'\s*\n*\s*\(NOTE.*\)*', '', cells[-1].text.encode('ascii', 'ignore')))    
    
    #if ie_type.find('Usage Report') != -1:
    if ie_type == 'Usage Report':
        if comment.find('Report Type') != -1:
            ie_type = "Usage Report Session Report Request"
        elif comment.find('Query URR') != -1:
            ie_type = "Usage Report Session Modification Response"
        elif comment.find('provisioned ') != -1:
            ie_type = "Usage Report Session Deletion Response"
        else:
             assert False, "Unknown IE type : [Usage Report]"
    
    if ie_type == 'Update BAR':
        if comment.find('7.5.4.11-1') != -1:
            ie_type = "Update BAR Session Modification Request"
        elif comment.find('7.5.9.2-1') != -1:
            ie_type = "Update BAR PFCP Session Report Response"
        else:
             assert False, "Unknown IE type : [Update BAR]"
    
    if ie_type.find('PFD Contents') != -1:
        ie_type = 'PFD contents'
    elif ie_type.find('PFD') != -1:
        ie_type = 'PFD context'
    elif ie_type.find('UE IP address') != -1:
        ie_type = 'UE IP Address'
    elif ie_type.find('SxSMReq-Flags') != -1:
        ie_type = 'PFCPSMReq-Flags'
    elif ie_type.find('PFCPSRRsp-Flags2') != -1:
        ie_type = 'PFCPSRRsp-Flags'
    elif ie_type.find('IPv4 Configuration Parameters (IP4CP)') != -1:
        ie_type = 'IP4CP'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[-1].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text.encode('ascii', 'ignore')
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text).encode('ascii', 'ignore')
    if ie_value[len(ie_value)-1] == ' ':
        ie_value = ie_value[:len(ie_value)-1]

    if ie_type == 'Create PDR' or ie_type == 'Create FAR' or ie_type == 'Update PDR':
        instance = "1"

    if int(instance) > int(type_list[ie_type]["max_instance"]):
        type_list[ie_type]["max_instance"] = instance
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_instance\"] = \"" + instance + "\"\n")

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "instance" : instance, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"instance\" : \"" + cells["instance"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cache = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r') 
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv-msg-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    msg_table = ""
    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            continue;
        else:
            if cell.text.find('Message Type value') != -1:
                msg_table = table
                d_print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-3]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        type = row.cells[0].text.encode('ascii', 'ignore')
        if type.isdigit() is False:
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv-type-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    ie_table = ""
    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            pass
        else:
            if cell.text.find('IE Type value') != -1:
                ie_table = table
                d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:-1]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\(', '', key)
        key = re.sub('\)', '', key)
        key = re.sub('\s*$', '', key)

        type = row.cells[0].text.encode('ascii', 'ignore')
        type_list[key] = { "type": type , "max_instance" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_instance\" : \"0\" }\n")
    f.close()

d_info("[Group IE List]")
cachefile = cachedir + 'tlv-group-list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    for i, table in enumerate(document.tables):
        try:
            cell = table.rows[0].cells[0]
        except:
            pass
        else:
            if cell.text.find('Octet') != -1 and \
               table.rows[0].cells[1].text.find('Outer Header to be created') == -1:

                num = 0;
                if len(table.rows[0].cells) > 2 and table.rows[0].cells[2].text.find('IE Type') != -1:
                    num = 2
                elif len(table.rows[0].cells) > 3 and table.rows[0].cells[3].text.find('IE Type') != -1:
                    num = 3
                elif len(table.rows[0].cells) > 4 and table.rows[0].cells[4].text.find('IE Type') != -1:
                    num = 4

                if num == 0:
                    continue;

                row = table.rows[0];

                d_print("Table Index = %d[%s]\n" % (i, row.cells[num].text))

                if len(re.findall('\d+', row.cells[num].text)) == 0:
                    continue;
                ie_type = re.findall('\d+', row.cells[num].text)[-1].encode('ascii', 'ignore')
                ie_name = re.sub('\s*IE Type.*', '', row.cells[num].text.encode('ascii', 'ignore'))

                d_print("TYPE:%s NAME:%s\n" % (ie_type, ie_name))

                # SKIP Access Forwarding Action Information
                if (int(ie_type) == 78):
                    ie_name =  "Usage Report Session Modification Response"
                elif (int(ie_type) == 79):
                    ie_name =  "Usage Report Session Deletion Response"
                elif (int(ie_type) == 80):
                    ie_name =  "Usage Report Session Report Request"    
                elif (int(ie_type) == 86):
                    ie_name =  "Update BAR Session Modification Request" 
                elif (int(ie_type) == 12):
                    ie_name =  "Update BAR PFCP Session Report Response" 

                if ie_name.find('Access Forwarding Action Information 2') != -1:
                    ie_idx = str(int(ie_type)+100)
                    write_file(f, "ies = []\n")
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
                    continue
                
                if ie_name not in group_list.keys():
                    ies = []
                    write_file(f, "ies = []\n")
                    for row in table.rows[4:]:
                        cells = get_cells(row.cells)
                        if cells is None:
                            continue

                        ies_is_added = True
                        for ie in ies:
                            if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                                ies_is_added = False
                        if ies_is_added is True:
                            ies.append(cells)
                            write_cells_to_file("ies", cells)

                    ie_idx = str(int(ie_type)+100)
                    group_list[ie_name] = { "index" : ie_idx, "type" : ie_type, "ies" : ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"index\" : \"" + ie_idx + "\", \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
    f.close()

msg_list["PFCP Heartbeat Request"]["table"] = 7
msg_list["PFCP Heartbeat Response"]["table"] = 8 
msg_list["PFCP PFD Management Request"]["table"] = 9
msg_list["PFCP PFD Management Response"]["table"] = 12
msg_list["PFCP Association Setup Request"]["table"] = 13
msg_list["PFCP Association Setup Response"]["table"] = 14
msg_list["PFCP Association Update Request"]["table"] = 15
msg_list["PFCP Association Update Response"]["table"] = 16
msg_list["PFCP Association Release Request"]["table"] = 17
msg_list["PFCP Association Release Response"]["table"] = 18
msg_list["PFCP Version Not Supported Response"]["table"] = 0
msg_list["PFCP Node Report Request"]["table"] = 19
msg_list["PFCP Node Report Response"]["table"] = 21
msg_list["PFCP Session Set Deletion Request"]["table"] = 22
msg_list["PFCP Session Set Deletion Response"]["table"] = 23
msg_list["PFCP Session Establishment Request"]["table"] = 24
msg_list["PFCP Session Establishment Response"]["table"] = 40
msg_list["PFCP Session Modification Request"]["table"] = 45
msg_list["PFCP Session Modification Response"]["table"] = 65
msg_list["PFCP Session Deletion Request"]["table"] = 67
msg_list["PFCP Session Deletion Response"]["table"] = 68
msg_list["PFCP Session Report Request"]["table"] = 70
msg_list["PFCP Session Report Response"]["table"] = 76

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "tlv-msg-" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            execfile(cachefile)
            print "Read from " + cachefile
        else:
            document = Document(filename)
            f = open(cachefile, 'w') 

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            if key.find('Association') != -1:
                start_i = 1
            elif key.find('Heartbeat') != -1:
                start_i = 1
            else:
                start_i = 2
            
            if key != "PFCP Session Deletion Request" and key != "PFCP Version Not Supported Response":
                for row in table.rows[start_i:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue
    
                    if (cells["ie_type"] == 'Create PDR' or cells["ie_type"] == 'Create FAR' or cells["ie_type"] == 'Update PDR'):
                        cells["instance"] = '0' 
                        cells["presence"] = 'O'
                        ies.append(cells)
                        write_cells_to_file("ies", cells)
                    cells = get_cells(row.cells)

                    ies_is_added = True
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        ies.append(cells)
                        write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

#type_list["Recovery"]["size"] = 1                       # Type : 3
#type_list["EBI"]["size"] = 1                            # Type : 73
#type_list["RAT Type"]["size"] = 1                       # Type : 82
#type_list["PDN Type"]["size"] = 1                       # Type : 99
#type_list["Port Number"]["size"] = 2                    # Type : 126
#type_list["APN Restriction"]["size"] = 1                # Type : 127
#type_list["Selection Mode"]["size"] = 1                 # Type : 128
#type_list["Node Type"]["size"] = 1                 # Type : 128

f = open(outdir + 'message.h', 'w')
output_header_to_file(f)
f.write("""#if !defined(OGS_PFCP_INSIDE) && !defined(OGS_PFCP_COMPILATION)
#error "This header cannot be included directly."
#endif

#ifndef OGS_PFCP_MESSAGE_H
#define OGS_PFCP_MESSAGE_H

#ifdef __cplusplus
extern "C" {
#endif

/* 5.1 General format */
#define OGS_PFCP_HEADER_LEN 16
#define OGS_PFCP_SEID_LEN   8
typedef struct ogs_pfcp_header_s {
    union {
        struct {
        ED4(uint8_t version:3;,
            uint8_t spare1:3;,
            uint8_t mp:1;,
            uint8_t seid_p:1;)
        };
        uint8_t flags;
    };
    uint8_t type;
    uint16_t length;
    union {
        struct {
            uint64_t seid;
            /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
#define OGS_PFCP_XID_TO_SQN(__xid) htonl(((__xid) << 8))
#define OGS_PFCP_SQN_TO_XID(__sqn) (ntohl(__sqn) >> 8)
            uint32_t sqn;
        };
        /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
        uint32_t sqn_only;
    };
} __attribute__ ((packed)) ogs_pfcp_header_t;

/* PFCP message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define OGS_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define OGS_PFCP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        if v_lower(k)=="cause" or v_lower(k)=="sequence_number" or v_lower(k)=="f_teid":
            f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k))
        else:
            f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

for k, v in group_list.items():
    if v_lower(k) == "ethernet_packet_filter":
        v["index"] = "1"
    if v_lower(k) == "pdi":
        v["index"] = "2"
    if v_lower(k) == "create_pdr":
        v["index"] = "3"
    if v_lower(k) == "forwarding_parameters":
        v["index"] = "4"
    if v_lower(k) == "duplicating_parameters":
        v["index"] = "5"
    if v_lower(k) == "create_far":
        v["index"] = "6"
    if v_lower(k) == "update_forwarding_parameters":
        v["index"] = "7"
    if v_lower(k) == "update_duplicating_parameters":
        v["index"] = "8"
    if v_lower(k) == "update_far":
        v["index"] = "9"
    if v_lower(k) == "pfd_context":
        v["index"] = "10"
    if v_lower(k) == "application_id_s_pfds":
        v["index"] = "11"
    if v_lower(k) == "ethernet_traffic_information":
        v["index"] = "12"
    if v_lower(k) == "access_forwarding_action_information_1":
        v["index"] = "13"
    if v_lower(k) == "access_forwarding_action_information_2":
        v["index"] = "14"
    if v_lower(k) == "update_access_forwarding_action_information_1":
        v["index"] = "15"
    if v_lower(k) == "update_access_forwarding_action_information_2":
        v["index"] = "16"

tmp = [(k, v["index"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]), reverse=False)

f.write("/* Group Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern ogs_tlv_desc_t ogs_pfcp_tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Infomration Element */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("typedef ogs_tlv_uint8_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 2:
            f.write("typedef ogs_tlv_uint16_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 3:
            f.write("typedef ogs_tlv_uint24_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 4:
            f.write("typedef ogs_tlv_uint32_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef ogs_tlv_octet_t ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

tmp = []
f.write("/* Structure for Group Infomration Element */\n")
for (k, v) in sorted_group_list:
    f.write("typedef struct ogs_pfcp_tlv_" + v_lower(k) + "_s {\n")
    f.write("    ogs_tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                v_lower(ies["ie_value"]))
        if ies["ie_type"] == "F-TEID":
            if ies["ie_value"] == "S2b-U ePDG F-TEID":
                f.write("_" + ies["instance"] + ";")
            elif ies["ie_value"] == "S2a-U TWAN F-TEID":
                f.write("_" + ies["instance"] + ";")
            else:
                f.write(";")
            f.write(" /* Instance : " + ies["instance"] + " */\n")
        else:
            f.write(";\n")
    f.write("} ogs_pfcp_tlv_" + v_lower(k) + "_t;\n")
    f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct ogs_" + v_lower(k) + "_s {\n")
        for ies in msg_list[k]["ies"]:
            # 0403 modify
            if ies["instance"] != "0":
                f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ies["instance"] + ";\n")
            else:
                f.write("    ogs_pfcp_tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
        f.write("} ogs_" + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct ogs_pfcp_message_s {\n")
f.write("   ogs_pfcp_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        ogs_" + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} ogs_pfcp_message_t;\n\n")

f.write("""int ogs_pfcp_parse_msg(ogs_pfcp_message_t *pfcp_message, ogs_pkbuf_t *pkbuf);
ogs_pkbuf_t *ogs_pfcp_build_msg(ogs_pfcp_message_t *pfcp_message);

#ifdef __cplusplus
}
#endif

#endif /* OGS_PFCP_MESSAGE_H */
""")
f.close()

f = open(outdir + 'message.c', 'w')
output_header_to_file(f)
f.write("""#include "ogs-pfcp.h"

""")

for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        if v_lower(k)=="cause" or v_lower(k)=="sequence_number" or v_lower(k)=="f_teid":
            f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        else:
            f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        if "size" in type_list[k]:
            if type_list[k]["size"] == 1:
                f.write("    OGS_TLV_UINT8,\n")
            elif type_list[k]["size"] == 2:
                f.write("    OGS_TLV_UINT16,\n")
            elif type_list[k]["size"] == 3:
                f.write("    OGS_TLV_UINT24,\n")
            elif type_list[k]["size"] == 4:
                f.write("    OGS_TLV_UINT32,\n")
            else:
                assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
        else:
            f.write("    OGS_TLV_VAR_STR,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    OGS_PFCP_%s_TYPE,\n" % v_upper(k))
        if "size" in type_list[k]:
            f.write("    %d,\n" % type_list[k]["size"])
        else:
            f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(ogs_pfcp_tlv_%s_t),\n" % v_lower(k))
        f.write("    { NULL }\n")
        f.write("};\n\n")

for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        f.write("    OGS_TLV_COMPOUND,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    OGS_PFCP_%s_TYPE,\n" % v_upper(k))
        f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(ogs_pfcp_tlv_%s_t),\n" % v_lower(k))
        f.write("    {\n")
        for ies in group_list[k]["ies"]:
                if v_lower(ies["ie_type"])=="cause" or v_lower(ies["ie_type"])=="sequence_number" or v_lower(ies["ie_type"])=="f_teid":
                    f.write("        &ogs_pfcp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
                else:
                    f.write("        &ogs_pfcp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("        NULL,\n")
        f.write("    }\n")
        f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("ogs_tlv_desc_t ogs_pfcp_tlv_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    OGS_TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
                if v_lower(ies["ie_type"])=="cause" or v_lower(ies["ie_type"])=="sequence_number" or v_lower(ies["ie_type"])=="f_teid":
                    f.write("        &ogs_pfcp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
                else:
                    f.write("        &ogs_pfcp_tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""int ogs_pfcp_parse_msg(ogs_pfcp_message_t *pfcp_message, ogs_pkbuf_t *pkbuf)
{
    int rv = OGS_ERROR;
    ogs_pfcp_header_t *h = NULL;
    uint16_t size = 0;

    ogs_assert(pfcp_message);
    ogs_assert(pkbuf);
    ogs_assert(pkbuf->len);

    h = (ogs_pfcp_header_t *)pkbuf->data;
    ogs_assert(h);
    
    memset(pfcp_message, 0, sizeof(ogs_pfcp_message_t));

    if (h->seid_p)
        size = OGS_PFCP_HEADER_LEN;
    else
        size = OGS_PFCP_HEADER_LEN-OGS_PFCP_SEID_LEN;

    ogs_assert(ogs_pkbuf_pull(pkbuf, size));
    memcpy(&pfcp_message->h, pkbuf->data - size, size);

    if (h->seid_p) {
        pfcp_message->h.seid = be64toh(pfcp_message->h.seid);
    } else {
        pfcp_message->h.sqn = pfcp_message->h.sqn_only;
        pfcp_message->h.sqn_only = pfcp_message->h.sqn_only;
    }

    if (pkbuf->len == 0)
        return OGS_OK;

    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case OGS_%s_TYPE:\n" % v_upper(k))
        f.write("            rv = ogs_tlv_parse_msg(&pfcp_message->%s,\n" % v_lower(k))
        f.write("                    &ogs_pfcp_tlv_desc_%s, pkbuf, OGS_TLV_MODE_T2_L2);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            ogs_warn("Not implmeneted(type:%d)", pfcp_message->h.type);
            break;
    }

    return rv;
}

""")

f.write("""ogs_pkbuf_t *ogs_pfcp_build_msg(ogs_pfcp_message_t *pfcp_message)
{
    ogs_pkbuf_t *pkbuf = NULL;

    ogs_assert(pfcp_message);
    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case OGS_%s_TYPE:\n" % v_upper(k))
        f.write("            pkbuf = ogs_tlv_build_msg(&ogs_pfcp_tlv_desc_%s,\n" % v_lower(k))
        f.write("                    &pfcp_message->%s, OGS_TLV_MODE_T2_L2);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            ogs_warn("Not implmeneted(type:%d)", pfcp_message->h.type);
            break;
    }

    return pkbuf;
}
""")

f.write("\n")

f.close()
